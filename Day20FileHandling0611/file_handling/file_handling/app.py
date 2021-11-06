from docx import Document #thao tác mở nó lên ( nó: wordoffice ?)
from docx.shared import Cm

doc = Document()#=new: tạo mới documnet
# doc = docx.document.Document
# doc.add_picture('avatar.jpg', width=Cm(4))
doc.add_paragraph("Ba Nguyễn")
doc.add_paragraph("Mô tả bản thân...")
doc.add_paragraph("Thông tin liên hệ...")
doc.add_heading('Thông tin bản thân...', level=2)
para = doc.add_paragraph("Tui là ")
para.add_run("Ba").bold = True
para.add_run(". Tui đến từ ")
para.add_run("Nam Định").italic = True

experiences = (
    ('Developer', 'Developer....'),
    ('Teacher', 'Teacher....'),
)

doc.add_heading("Kinh nghiệm làm việc...", level=2)
table = doc.add_table(rows=1, cols=2)
position, detail = table.rows[0].cells

position.text = 'Vị trí'
detail.text = 'Chi tiết'

for p, d in experiences:
    position, detail = table.add_row().cells
    position.text = p
    detail.text = d

doc.save('cv.docx')