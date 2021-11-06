from docx import Document
def xin_tang_luong(ho_ten, luong_hien_tai):
    doc = Document()
    doc.add_heading('Don xin tang luong', level=0)
    doc.add_paragraph(f'Xin chao chi quan ly, em ten la {ho_ten}, em muon tang luong, vi luong hien tai {luong_hien_tai} nhu nay la qua thap')
    doc.save('don_xin_tang_luong.docx')

if __name__== "__main__":
    ho_ten = input("Ho ten: ")
    luong_hien_tai = input("Luong hien tai: ")
    xin_tang_luong(ho_ten, luong_hien_tai)